import os
from docx import Document
import pytest
from src.ingestion.chunking import process_markdown_for_parent_child
from src.ingestion.docx_utils import convert_docx_to_markdown

def test_convert_docx_to_markdown(tmp_path):
    docx_path = str(tmp_path / "test.docx")
    doc = Document()
    
    # Add Headings
    doc.add_heading("Main Title", 0) # Title
    doc.add_heading("Section 1", 1) # Heading 1
    doc.add_paragraph("This is a normal paragraph.")
    
    # Add Bullet list
    doc.add_paragraph("First bullet", style="List Bullet")
    doc.add_paragraph("Second bullet", style="List Bullet")
    
    # Add Numbered list
    doc.add_paragraph("First number", style="List Number")
    
    # Add Table
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Data 1"
    table.cell(1, 1).text = "Data 2"
    
    doc.save(docx_path)
    
    md_text = convert_docx_to_markdown(docx_path)
    
    # Verify outputs
    assert "# Main Title" in md_text
    assert "# Section 1" in md_text
    assert "This is a normal paragraph." in md_text
    assert "- First bullet" in md_text
    assert "1. First number" in md_text
    assert "| Header 1 | Header 2 |" in md_text
    assert "|---|---|" in md_text
    assert "| Data 1 | Data 2 |" in md_text


def test_docx_parent_child_routing():
    md_content = """
# Main Title
Intro text

## Section 1
Detailed content
"""
    parent_dict, child_chunks = process_markdown_for_parent_child(md_content, "test.docx", "main", doc_type="document_markdown")
    
    # Check parent_dict
    assert len(parent_dict) == 2
    titles = [p["topic_title"] for p in parent_dict.values()]
    assert "Main Title" in titles
    assert "Section 1" in titles
    
    # Check types
    assert all(p["type"] == "document_markdown" for p in parent_dict.values())
    assert all(c["type"] == "document_markdown" for c in child_chunks)
    assert all(c["metadata"]["type"] == "document_markdown" for c in child_chunks)
