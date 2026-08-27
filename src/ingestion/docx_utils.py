"""
src/ingestion/docx_utils.py
============================
DOCX → Markdown conversion utilities for the SnowWiki ingestion pipeline.

Separates DOCX-specific conversion logic from the core chunking algorithms
in chunking.py so each module stays focused on a single responsibility.

Public API
----------
convert_docx_to_markdown(file_path: str) -> str
"""

from __future__ import annotations

import re
import logging

from docx import Document

logger = logging.getLogger(__name__)


# ── Internal helpers ───────────────────────────────────────────────────────────

def _is_servicenow_code(text: str) -> bool:
    """
    Weighted pattern-matching heuristic for ServiceNow-specific script detection.
    Require at least 1 High Confidence Trigger OR 2 Syntax Multi-Matches.
    """
    high_confidence = [
        r"g_form\.", r"gs\.", r"new GlideRecord", r"GlideAggregate",
        r"current\.setValue", r"action\.setRedirectURL", r"JSON\.parse"
    ]
    multi_match = [
        r"function\s+\w+\s*\(", r"\}\s*;", r"return\s+false;"
    ]

    hc_matches = sum(1 for p in high_confidence if re.search(p, text))
    if hc_matches >= 1:
        return True

    mm_matches = sum(1 for p in multi_match if re.search(p, text))
    if mm_matches >= 2:
        return True

    return False

def _sanitize_table_cells(md_text: str) -> str:
    """
    Multi-Line Table Cell Sanitization: Replace internal \n in table cells with <br>.
    Mammoth output creates tables where cell contents might contain newlines.
    """
    lines = md_text.splitlines()
    for i, line in enumerate(lines):
        if line.strip().startswith('|') and line.strip().endswith('|'):
            # Replace actual newlines inside the row with <br>
            lines[i] = line.replace("\\n", "<br>")
    return "\n".join(lines)


def _normalize_docx_markdown(md_text: str) -> str:
    """Enhance and fix Markdown generated from DOCX."""
    # First, sanitize table cells
    md_text = _sanitize_table_cells(md_text)

    lines = md_text.splitlines()
    out_lines = []

    topic_pattern    = re.compile(r"^(Topic\s+\d+[:\-]?.*)$", re.IGNORECASE)
    numbered_pattern = re.compile(r"^(\d+\.\s+[A-Z][a-zA-Z0-9\s&]+)$")

    in_code_block = False

    for line in lines:
        line_clean = line.strip()
        
        # Toggle code block state
        if line_clean.startswith("```"):
            in_code_block = not in_code_block
            out_lines.append(line)
            continue
            
        if in_code_block:
            # Context-Aware Un-escaping: bypass code fence blocks
            out_lines.append(line)
            continue

        if not line_clean:
            out_lines.append("")
            continue

        # Context-aware unescaping for prose blocks (Fix mammoth escapes)
        line = line.replace("\\[", "[").replace("\\]", "]")
        line = line.replace("\\$", "$").replace("\\.", ".")
        line = line.replace("\\_", "_")

        # Heading Hierarchy Normalization
        if line_clean.startswith("**") and line_clean.endswith("**"):
            inner = line_clean.strip("*").strip()
            if topic_pattern.match(inner):
                out_lines.append(f"## {inner}")
                continue
            if numbered_pattern.match(inner):
                out_lines.append(f"### {inner}")
                continue
            if len(inner) < 80:
                out_lines.append(f"### {inner}")
                continue

        if topic_pattern.match(line_clean):
            out_lines.append(f"## {line_clean}")
            continue

        if numbered_pattern.match(line_clean) and not line_clean.startswith("#"):
            out_lines.append(f"### {line_clean}")
            continue

        line = line.replace("\xa0", " ")
        out_lines.append(line)

    # Secondary pass to heuristically wrap code blocks if not already fenced
    final_text = "\n".join(out_lines)
    
    paras = final_text.split("\n\n")
    for i, p in enumerate(paras):
        if not p.strip().startswith("```") and _is_servicenow_code(p):
            # Wrap in javascript fences
            paras[i] = f"```javascript\n{p.strip()}\n```"
            
    return "\n\n".join(paras)


def _fallback_convert_docx_to_markdown(file_path: str) -> str:
    """Pure-python DOCX → Markdown conversion used when mammoth is unavailable."""
    doc = Document(file_path)
    md_lines = []

    topic_pattern    = re.compile(r"^(Topic\s+\d+[:\-]?.*)$", re.IGNORECASE)
    numbered_pattern = re.compile(r"^(\d+\.\s+[A-Z][a-zA-Z0-9\s&]+)$")

    for p in doc.paragraphs:
        style_name = p.style.name if p.style else ""
        text = p.text.strip()
        if not text:
            continue

        is_bold_run = any(run.bold for run in p.runs)

        if style_name.startswith("Heading"):
            try:
                level = int(style_name.split()[-1])
                md_lines.append(f"{'#' * level} {text}\n")
            except ValueError:
                md_lines.append(f"# {text}\n")
        elif style_name == "Title":
            md_lines.append(f"# {text}\n")
        elif topic_pattern.match(text) or numbered_pattern.match(text):
            md_lines.append(f"## {text}\n")
        elif is_bold_run and len(text) < 80:
            md_lines.append(f"### {text}\n")
        elif style_name.startswith("List Bullet") or style_name == "List Paragraph":
            md_lines.append(f"- {text}")
        elif style_name.startswith("List Number"):
            md_lines.append(f"1. {text}")
        else:
            if _is_servicenow_code(text):
                md_lines.append(f"```javascript\n{text}\n```\n")
            else:
                md_lines.append(text + "\n")

    for table in doc.tables:
        if not table.rows:
            continue
        header = table.rows[0]
        header_texts = [cell.text.strip().replace("\n", "<br>") for cell in header.cells]
        md_lines.append("\n| " + " | ".join(header_texts) + " |")
        md_lines.append("|" + "|".join(["---"] * len(header.cells)) + "|")
        for row in table.rows[1:]:
            row_texts = [cell.text.strip().replace("\n", "<br>") for cell in row.cells]
            md_lines.append("| " + " | ".join(row_texts) + " |")
        md_lines.append("\n")

    return "\n".join(md_lines)


# ── Public API ─────────────────────────────────────────────────────────────────

def convert_docx_to_markdown(file_path: str) -> str:
    """
    Convert a DOCX file to structured Markdown using Mammoth.

    Translates Heading styles to Markdown headers (#, ##), bullets to lists,
    and applies semantic normalisation for pseudo-headings.
    Falls back to a pure-python implementation if mammoth is not installed.
    """
    try:
        import mammoth
        # Define style mapping for code blocks
        style_map = (
            "p[style-name='Code'] => pre:fresh\n"
            "p[style-name='HTML Preformatted'] => pre:fresh\n"
            "p[style-name='Source Code'] => pre:fresh"
        )
        with open(file_path, "rb") as docx_file:
            result = mammoth.convert_to_markdown(docx_file, style_map=style_map)
            raw_md = result.value
            return _normalize_docx_markdown(raw_md)
    except ImportError:
        logger.warning("mammoth package not found. Using fallback DOCX to Markdown converter.")
        return _fallback_convert_docx_to_markdown(file_path)
